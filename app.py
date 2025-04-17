from flask import Flask, render_template, request, send_file, jsonify
from docx import Document
import os
from datetime import datetime
import json

app = Flask(__name__)

# Daily visitor count file path
visitor_count_file = "visitor_count.json"

# Initialize visitor count for the day
def initialize_visitor_count():
    if not os.path.exists(visitor_count_file):
        with open(visitor_count_file, "w") as f:
            json.dump({"date": str(datetime.now().date()), "count": 0}, f)

def increment_visitor_count():
    # Read the current visitor count
    with open(visitor_count_file, "r") as f:
        data = json.load(f)
    
    # Check if the date is still the same
    current_date = str(datetime.now().date())
    if data["date"] == current_date:
        data["count"] += 1  # Increment the count for today
    else:
        data["date"] = current_date
        data["count"] = 1  # Reset the count for the new day
    
    # Save the updated count
    with open(visitor_count_file, "w") as f:
        json.dump(data, f)

    return data["count"]

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        try:
            # Collect form data for both Assignment and Lab Report
            form_data = {}

            cover_type = request.form["cover_type"]

            # Initialize visitor count for the day
            visitor_count = increment_visitor_count()

            # Collect form data for Assignment
            if cover_type == "Assignment":
                form_data = {
                    "No": request.form.get("assignment_no", ""),  # Optional
                    "Here_Course_Code": request.form["course_code"],
                    "Here_Course_Title": request.form["course_title"],
                    "Here_TeacherName": request.form["teacher_name"],
                    "teacher_designation": request.form["teacher_designation"],
                    "Here_Teachers_Department_Name": request.form["teacher_dept"],
                    "Here_StudentName": request.form["student_name"],
                    "Here_StudentID": request.form["student_id"],
                    "Here_Section": request.form["student_section"],
                    "Here_DepartmentNameOfStudent": request.form["department_name"],
                    "HereDate": request.form["submission_date"]
                }
                doc = Document("Assignment Cover Page.docx")
            else:
                # Collect form data for Lab Report
                form_data = {
                    "Course_Title": request.form["course_title"],
                    "Course_Code": request.form["course_code"],
                    "designation": request.form["teacher_designation"],
                    "TeacherDeptName": request.form["teacher_dept"],
                    "StudentName": request.form["student_name"],
                    "TeacherName": request.form["teacher_name"],
                    "StudentID": request.form["student_id"],
                    "Group": request.form["student_section"],
                    "DPT": request.form["department_name"],
                    "DateIs": request.form["submission_date"],  # Date for Lab Report
                    "HereExperimentNo": request.form["experiment_no"],  # Experiment No
                    "HereExperimentName": request.form["experiment_name"]  # Experiment Name
                }
                doc = Document("LabReport Template.docx")

            # Loop through the paragraphs and replace placeholders
            for para in doc.paragraphs:
                for key, value in form_data.items():
                    if key in para.text:
                        for run in para.runs:
                            if key in run.text:
                                run.text = run.text.replace(key, value)

            # Save the generated DOCX file
            os.makedirs("output", exist_ok=True)
            date_str = datetime.now().strftime("%Y-%m-%d")
            docx_path = os.path.join("output", f"CoverPage_{date_str}.docx")
            doc.save(docx_path)

            # Return success response with file path
            return jsonify({"success": True, "file": docx_path, "visitor_count": visitor_count})

        except Exception as e:
            return jsonify({"success": False, "error": str(e)})

    # For GET request, show the form and current visitor count
    return render_template("form.html", visitor_count=increment_visitor_count())

@app.route("/download")
def download():
    path = request.args.get("file")
    return send_file(path, as_attachment=True)

if __name__ == "__main__":
    # Initialize visitor count when the app starts
    initialize_visitor_count()
    app.run(debug=True)

